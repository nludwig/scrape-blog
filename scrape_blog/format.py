import logging

from docx import Document

def format_paragraphs_to_docx(title, paras, doc=None, super_title=None):
    if doc is None:
        doc = Document()
        if super_title is not None:
            doc.add_heading(super_title, 0)
            doc.add_page_break()

    doc.add_heading(title)

    for para in paras:
        if isinstance(para, dict):
            logging.info('Got image of width %s', para.get('width'))
            try:
                doc.add_picture(para['img'], width=para.get('width'))
            except Exception:
                logging.error(
                        'Exception while adding picture:',
                        exc_info=True,
                )
            finally:
                try:
                    para['img'].close()
                except Exception:
                    logging.error(
                            'Exception while closing picture:',
                            exc_info=True,
                    )
        else:
            doc.add_paragraph(para)

    doc.add_page_break()

    return doc
